import docx
from os import path
from .hyperlink import add

def createDocxIndex(files, folder, attorney):
    index = docx.Document()
    index.add_heading('Indice atti e documenti')

    for file in files:
        
        
        p = index.add_paragraph()
        p.style = 'List Number 2'
        try:
            new_run = p.add_run(file.name)
            add(p, new_run, file.url)
        except ValueError:
            print('Error: '+ file.name)

    index.add_paragraph(f"Avv. {attorney}")
    

    index.save(path.join(folder,'00_indice_atti_documenti.docx'))
